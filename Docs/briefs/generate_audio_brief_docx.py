"""
One-shot generator for the SolShot audio brief Word doc.

Run with: python Docs/briefs/generate_audio_brief_docx.py
Outputs:  Docs/briefs/SolShot_Audio_Brief.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).parent / "SolShot_Audio_Brief.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

ACCENT = RGBColor(0xC8, 0x78, 0x1A)  # SolShot accent orange
OLIVE = RGBColor(0x7A, 0x90, 0x60)
DEEP = RGBColor(0x2A, 0x33, 0x1F)

def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DEEP
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)

def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DEEP
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(text, indent=0):
    p = doc.add_paragraph(text, style="List Bullet")
    if indent:
        p.paragraph_format.left_indent = Inches(0.25 * indent)

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = DEEP
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)

def hr():
    p = doc.add_paragraph()
    r = p.add_run("─" * 80)
    r.font.color.rgb = OLIVE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

# ─── HEADER ────────────────────────────────────────────────
h1("Sound Design Brief: SolShot")
p = doc.add_paragraph()
r = p.add_run("Audio identity for a Solana-based artillery PvP game in the iShoot / Pocket Tanks lineage.")
r.italic = True
r.font.color.rgb = OLIVE
hr()

# ─── CONTEXT ───────────────────────────────────────────────
h2("Context")
body(
    "SolShot is a 1v1 + multi-player artillery PvP game on Solana, shipped as a web app "
    "(solshot.gg) and a Telegram Mini App. Think Pocket Tanks / iShoot / Worms lineage — "
    "turn-based, side-view, projectile arcs, destructible terrain. Spiritual reference is "
    "iShoot (2008–09 iPhone) — modest audio that was iconic, not Hollywood. We want SolShot "
    "to land in that space."
)
body(
    "Game has 20 weapons across 7 tiers, 6 terrain biomes, two match modes (real-time 1v1 / "
    "async multi-player FFA in Telegram groups), HP-based damage with weapon-specific blast radii."
)

# ─── WHAT I NEED ───────────────────────────────────────────
h2("What I Need")
body(
    "A coherent audio identity — every weapon recognisable by its fire+impact alone, every UI "
    "moment with a satisfying click, six biome ambiences that don't compete with action SFX. "
    "Game audio at the \"polished iShoot for 2026\" bar — not \"AAA console cinematic\". "
    "Single dev / single composer scope."
)

# ─── TONE REFERENCE ────────────────────────────────────────
h2("Tone Reference")

h3("Yes")
bullet("iShoot weapon variety — each shot reads as itself")
bullet("Pocket Tanks restraint — no overproduced layers")
bullet("Hi-Fi Rush music-as-feedback — every UI tick has rhythm")
bullet("Worms Mobile impact satisfaction — the \"thunk\" matters more than the boom")
bullet("Lo-fi pixel-art shading on the SFX — slightly compressed, slightly crunchy, like 90s arcade")

h3("No")
bullet("Hollywood blockbuster bass drops")
bullet("Cinematic risers / drones")
bullet("Generic Unreal Engine Marketplace SFX packs")
bullet("Mobile-game over-mixed UI clicks (Candy Crush territory)")
bullet("Vocal stings (\"HEADSHOT!\" / \"DOUBLE KILL!\")")
bullet("Anything that requires sub-bass to read on a phone speaker")

# ─── DELIVERABLES ──────────────────────────────────────────
h2("Deliverables")

# 1
h3("1 — UI Sound Set (~25 cues, 5–10 min total)")

body("Buttons / navigation:")
bullet("ui_tap — primary button press (PLAY, FIRE, etc.)")
bullet("ui_tap_secondary — secondary button (subtle)")
bullet("ui_hover — desktop hover (optional, very subtle)")
bullet("ui_back — back button / nav-back")
bullet("ui_screen_transition — between major screens (~300–500ms)")

body("Modals / toasts:")
bullet("ui_modal_open")
bullet("ui_modal_close")
bullet("ui_toast_success")
bullet("ui_toast_error")

body("Currency / counters:")
bullet("ui_gold_tick — for the +X gold popup, can be repeated")
bullet("ui_gold_earned_kill — bigger gold accent (kill bonus)")
bullet("ui_shot_burn — $SHOT burn for prestige (one-shot, ceremonial)")
bullet("ui_currency_change — wallet balance updates")

body("Match lifecycle:")
bullet("match_start_countdown_tick — 3, 2, 1")
bullet("match_start_fanfare — short, ~1.5s, deploy moment")
bullet("round_start / round_end — between-round transitions")
bullet("match_end_victory — ~3s, ceremonial but not bombastic")
bullet("match_end_defeat — ~2–3s, somber")
bullet("your_turn — turn-pass to local player")
bullet("opponent_turn — turn-pass to opponent (subtle)")
bullet("turn_timer_warning — repeating tick when ≤10s remain")

body("Eliminations:")
bullet("tank_eliminated — opponent KIA")
bullet("you_eliminated — local player KIA (stronger)")
bullet("buyback_purchase — group-chat re-entry (ceremonial)")

# 2
h3("2 — Combat Sound Set (~80 cues, the heart of the deliverable)")

body("For each of the 20 weapons, 3 cues:")
bullet("weap_<name>_fire — projectile launch (~200–400ms)")
bullet("weap_<name>_inflight — looping in-flight body (only some — see table)")
bullet("weap_<name>_impact — explosion / damage moment (~400–800ms)")

body("Weapon list with character notes:")

# Build a table for the weapon list
weapons = [
    ("0", "Single Shot", "FREE", "Crisp, neutral, recognizable. The default."),
    ("1", "Big Shot", "RARE", "Bigger boom, longer tail, satisfying."),
    ("2", "3 Shot", "TACTICAL", "3 quick crackle-pops at fire (one per projectile), 3 small impacts."),
    ("4", "Jackhammer", "EPIC", "Drill noise on impact, 5 chained smaller blasts."),
    ("5", "Heatseeker", "TACTICAL", "Whistle/whine on launch, swooping pitch in-flight, sharp impact."),
    ("7", "Pile Driver", "RARE", "Heavy mechanical thuds, 6 chained drill-impacts."),
    ("9", "Crazy Ivan", "LEGENDARY", "Chaos — 15 random small explosions in 1.5s, slightly comedic."),
    ("10", "Spider", "TACTICAL", "Mechanical click+spread on impact, 6 small leg-impact pops."),
    ("11", "Sniper Rifle", "RARE", "Sharp crack on fire, supersonic crack in flight, surgical 1px hit (no boom)."),
    ("12", "Magic Wall", "STANDARD", "Magical chime on impact, terrain-rises sound."),
    ("15", "Napalm", "RARE", "Whoosh on fire, sustained crackle on impact, ~2s burn."),
    ("16", "Hail Storm", "EPIC", "Whoosh on launch, rain-of-impacts (~10–15 pops over 1s)."),
    ("17", "Ground Hog", "EPIC", "Drill-down whir, muffled tunnel rumble, emerge+detonate boom."),
    ("20", "Skipper", "TACTICAL", "3–4 hopping bounces on terrain (each its own small thud)."),
    ("25", "Dirt Ball", "STANDARD", "Soft thud, terrain-rises sound (no damage)."),
    ("26", "Tommy Gun", "TACTICAL", "Rapid-fire 12-projectile burst, classic gangster vibe."),
    ("—", "5 Prestige weapons", "PRESTIGE", "Treat as flagship — burn-to-unlock rewards. Should feel premium."),
]

table = doc.add_table(rows=1 + len(weapons), cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, h in enumerate(["ID", "Name", "Tier", "Character"]):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.bold = True
for r, w in enumerate(weapons, start=1):
    row = table.rows[r].cells
    for i, val in enumerate(w):
        row[i].text = val

doc.add_paragraph()  # spacer

# 3
h3("3 — Damage Feedback (~10 cues)")
bullet("dmg_taken_self — you got hit")
bullet("dmg_dealt — you hit them (subtle accent so you know you connected)")
bullet("dmg_glancing — small damage (≤10 HP)")
bullet("dmg_solid — medium (10–50 HP)")
bullet("dmg_critical — heavy (50–100 HP)")
bullet("dmg_devastating — direct hit, near-KO (100+ HP)")
bullet("dmg_miss — projectile out of bounds, no impact")
bullet("dmg_self_damage — you damaged yourself (rare)")

# 4
h3("4 — Movement / Controls (~6 cues)")
bullet("tank_move_step — single tread \"chunk\", repeats while moving")
bullet("turret_rotate_tick — subtle servo, repeats while aiming")
bullet("power_adjust_tick — same family, slightly different pitch")
bullet("weapon_select — click between weapons in shop / arsenal")
bullet("weapon_owned_chime — when you successfully buy")

# 5
h3("5 — Ambient Biome Loops (~6 cues, 30–60s seamless)")
bullet("amb_jungle — insects, distant birds")
bullet("amb_arctic — sparse wind, ice creak")
bullet("amb_desert — dry wind, distant rumble")
bullet("amb_moon — near-silence, low subliminal rumble")
bullet("amb_volcanic — distant rumble, occasional sizzle")
bullet("amb_default — neutral / hub ambience")
body(
    "Keep these very low in the mix — they should sit under everything else and never compete "
    "with a weapon impact. Side-chain duck them if needed."
)

# 6
h3("6 — Music (1 main theme, optional)")
body(
    "Main menu theme — looping ~60–90s, sets the field-manual / military-radio tone. Sparse, "
    "not anthemic. Think the menu music of a 90s arcade artillery game, not a film score. "
    "Optional — can ship without and add later. If included, must duck under SFX, never compete."
)

# ─── FORMAT SPEC ────────────────────────────────────────────
h2("Format Spec")
bullet("Sample rate: 44.1kHz")
bullet("Bit depth: 16-bit")
bullet("Format: WAV masters + OGG/Vorbis exports")
bullet("Length: UI cues <300ms, combat fire/impact <800ms, ambient loops 30–60s seamless")
bullet("Mono unless stereo width matters — weapons can be stereo (especially scatter / multi-projectile), UI clicks should be mono")
bullet("Headroom: master peaks at -3dBTP, target loudness ~-14 LUFS for SFX (no aggressive limiter)")
bullet("Naming: snake_case, prefix-grouped — ui_tap.wav, weap_single_fire.wav, weap_single_impact.wav, dmg_critical.wav, amb_jungle.ogg")

# ─── CONSTRAINTS ────────────────────────────────────────────
h2("Constraints")
bullet("Mobile speaker compatibility — primary playback surface is a phone speaker through Telegram. Anything below ~100Hz won't reproduce. Mix accordingly; sub-bass is wasted.")
bullet("Polyphony budget — assume up to 8 SFX overlapping (4-player multi-shot scatter weapon during a kill). Cues should mix naturally, not all peak in the same frequency band.")
bullet("No vocal stings — no \"Double Kill!\", \"First Blood!\". Tone is silent-protagonist military, not arcade announcer.")
bullet("No copyrighted samples — must be original or licensed for game-commercial use, royalty-free in perpetuity, transferable.")

# ─── REFERENCES ─────────────────────────────────────────────
h2("Reference Tracks for Tone / Energy")
body("Anchor on the feel, don't have to literally match.")
bullet("iShoot (2008) — weapon variety, retro charm")
bullet("Pocket Tanks (2001) — UI restraint, weapon distinctiveness")
bullet("Worms Armageddon — slightly comedic violence energy (good for Crazy Ivan, Skipper)")
bullet("Hi-Fi Rush — UI rhythm and tactile clicks")
bullet("Hotline Miami — punchy lo-fi impact philosophy")

# ─── PROCESS ────────────────────────────────────────────────
h2("Process")
bullet("Round 1: 5 weapon sets (Single Shot, Heatseeker, Crazy Ivan, Sniper, Big Shot) + full UI cue set. We listen, decide if the tone is right.")
bullet("Round 2: if R1 lands, full deliverable.")
bullet("Round 3: revisions / mixing notes / stragglers.")

body("")
p = doc.add_paragraph()
r = p.add_run("Total budget: 80–120 cues + 6 ambient loops + (optional) 1 menu theme.")
r.bold = True
body("Estimated $800–1,500 on Fiverr/Upwork for an experienced single game-audio designer; $2,000–4,000 for a small studio with faster turnaround.")

# ─── OUTPUT STRUCTURE ───────────────────────────────────────
h2("Output Structure")
code("""solshot_audio_v1/
├── ui/
│   ├── ui_tap.wav
│   ├── ui_modal_open.wav
│   └── ...
├── weapons/
│   ├── weap_single_fire.wav
│   ├── weap_single_impact.wav
│   └── ...
├── damage/
│   └── dmg_critical.wav
├── ambient/
│   └── amb_jungle.ogg
├── music/
│   └── menu_theme.ogg
└── README.md  (mix notes, recommended ducking, polyphony tips)""")

# ─── CLOSING NOTES ──────────────────────────────────────────
h2("Notes for the Hiring Process")
bullet("Front-load Round 1 deliverable — 5 flagship weapons + UI set is enough to commit/reject the tone after a week. Don't pay for 80 cues before validating direction on 10.")
bullet("Reference tracks are the most important paragraph. \"iShoot for 2026\" is shorthand-magic for the right designer. Drop any of the references and the brief gets generic.")
bullet("Reject Hollywood mix on first listen. If R1 sounds like a movie trailer, it's wrong direction — pull back or change designer. iShoot was crunchy and direct, not cinematic.")

body("")
body("Good Fiverr search terms: \"retro arcade SFX\", \"indie game weapons sound\", \"8-bit modern hybrid\".")
body("Avoid: \"AAA cinematic SFX\", \"blockbuster game audio\".")

# Save
doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Size:  {OUT.stat().st_size:,} bytes")
