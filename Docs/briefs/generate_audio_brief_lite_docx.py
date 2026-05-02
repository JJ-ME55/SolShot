"""
Generates SolShot_Audio_Brief_Lite.docx — the trimmed "custom cues only"
hybrid brief for paste-into-Fiverr.

~27 cues, $400-600 budget. Pairs with SolShot_Audio_FreeLibrary_Shopping_List.md
which covers the ~54 commodity cues from free libraries.

Run with: python Docs/briefs/generate_audio_brief_lite_docx.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from pathlib import Path

OUT = Path(__file__).parent / "SolShot_Audio_Brief_Lite.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

ACCENT = RGBColor(0xC8, 0x78, 0x1A)
OLIVE = RGBColor(0x7A, 0x90, 0x60)
DEEP = RGBColor(0x2A, 0x33, 0x1F)


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DEEP
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = ACCENT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DEEP
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def hr():
    p = doc.add_paragraph()
    r = p.add_run("─" * 80)
    r.font.color.rgb = OLIVE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ─── HEADER ────────────────────────────────────────────────
h1("Sound Design Brief: SolShot — Custom Cues Only")
p = doc.add_paragraph()
r = p.add_run(
    "Hybrid scope: ~27 identity-defining cues for a Solana-based artillery PvP game. "
    "Long-tail commodity SFX (other weapons, generic UI, ambient loops) sourced from "
    "free libraries separately — please don't quote me for them."
)
r.italic = True
r.font.color.rgb = OLIVE
hr()

# ─── CONTEXT ───────────────────────────────────────────────
h2("Context")
body(
    "SolShot is a 1v1 + multi-player artillery PvP game on Solana, shipped as a web app "
    "(solshot.gg) and a Telegram Mini App. Think Pocket Tanks / iShoot / Worms lineage — "
    "turn-based, side-view, projectile arcs, destructible terrain. Spiritual reference is "
    "iShoot (2008–09 iPhone) — modest audio that was iconic, not Hollywood. We want SolShot "
    "to land in that space."
)

# ─── TONE REFERENCE ────────────────────────────────────────
h2("Tone Reference")

h3("Yes")
bullet("iShoot weapon variety — each shot reads as itself")
bullet("Pocket Tanks restraint — no overproduced layers")
bullet("Worms Mobile impact satisfaction — the \"thunk\" matters more than the boom")
bullet("Lo-fi pixel-art shading on the SFX — slightly compressed, slightly crunchy, like 90s arcade")

h3("No")
bullet("Hollywood blockbuster bass drops")
bullet("Cinematic risers / drones / orchestral hits")
bullet("Generic Unreal Engine Marketplace SFX packs")
bullet("Mobile-game over-mixed UI clicks (Candy Crush territory)")
bullet("Vocal stings (\"HEADSHOT!\" / \"DOUBLE KILL!\")")
bullet("Anything that requires sub-bass to read on a phone speaker")

# ─── DELIVERABLES ──────────────────────────────────────────
h2("Deliverables — 27 Custom Cues")

# Tier 1
h3("TIER 1 — Flagship Weapons (15 cues)")
body(
    "These are the weapons players hear thousands of times across a session. Mix-by-one-person "
    "consistency matters most here — each weapon must feel distinct, but they need to live in "
    "the same sonic world. Stems separated please (fire layer, impact layer) so we can layer "
    "them independently."
)

t1_table = doc.add_table(rows=7, cols=3)
t1_table.style = "Light Grid Accent 1"
t1_table.rows[0].cells[0].text = "Weapon"
t1_table.rows[0].cells[1].text = "Cues"
t1_table.rows[0].cells[2].text = "Character"
for cell in t1_table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.bold = True

weapons_t1 = [
    ("Single Shot", "fire + impact (2)", "Crisp, neutral, recognizable. The default. 80% of all shots fired use this."),
    ("Big Shot", "fire + impact (2)", "Bigger boom, longer tail, satisfying. The 'reward' weapon."),
    ("Heatseeker", "fire + inflight + impact (3)", "Whistle/whine on launch, swooping pitch in-flight, sharp impact."),
    ("Crazy Ivan", "fire + impact (2)", "Chaos. Impact gets repeated 15× in 1.5s for the scatter pattern. Slightly comedic."),
    ("Sniper Rifle", "fire + supersonic-inflight + minimal impact (3)", "Sharp crack on fire, supersonic crack in flight, surgical 1px hit (NO boom — that's the point)."),
    ("Ground Hog", "fire + tunnel-rumble + emerge-detonate (3)", "Drill-down whir, muffled tunnel rumble while underground, emerge+detonate boom."),
]
for i, (w, c, char) in enumerate(weapons_t1, start=1):
    t1_table.rows[i].cells[0].text = w
    t1_table.rows[i].cells[1].text = c
    t1_table.rows[i].cells[2].text = char

doc.add_paragraph()

# Tier 2
h3("TIER 2 — Identity / Match Lifecycle Moments (7 cues)")
body(
    "Emotional and branding cues. Stock would feel generic against the rest of the game's "
    "identity. These are short — most under 2s — but they're what players remember."
)

t2_items = [
    ("match_start_fanfare", "~1.5s. The first sound of every match. Deploy-moment energy. Crunchy, NOT bombastic."),
    ("match_end_victory", "~3s. Ceremonial win sting. Field-manual 'mission accomplished' vibe."),
    ("match_end_defeat", "~2-3s. Somber, not dramatic. A 'noted' — not a 'tragedy'."),
    ("you_eliminated", "Local-player KIA — the most emotionally-loaded cue in the game. Subtle 'final' weight."),
    ("tank_eliminated", "Opponent KIA — kill confirmation. Crisp, satisfying, NOT triumphant."),
    ("ui_shot_burn", "$SHOT prestige burn — fires once on rare ceremonial moments. Treat like sealing a wax stamp."),
    ("ui_gold_earned_kill", "Reward sting after kills (+200G bonus). Quick, punchy, gold-coin-ish but lo-fi."),
]
for name, desc in t2_items:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.name = "Consolas"
    p.add_run(f" — {desc}")

# Tier 3
h3("TIER 3 — Most-Clicked UI (5 cues)")
body(
    "The cues players hear hundreds of times per session. Mismatched stock here erodes "
    "identity fast. Keep these very short (under 100ms each, except dmg_critical/devastating)."
)

t3_items = [
    ("ui_tap", "Primary button press — FIRE button + every PLAY/CONFIRM/DEPLOY tap. Most-clicked sound in the app."),
    ("ui_modal_open / ui_modal_close", "Two paired cues. Page-transition vocabulary. Subtle, swift."),
    ("ui_toast_error", "Frequent. Sets the failure vocabulary. Short, slightly sour."),
    ("dmg_critical", "Heavy hit confirmation (50-100 HP damage band). The satisfying 'THUNK'."),
    ("dmg_devastating", "Direct hit, near-KO (100+ HP). Bigger and bassier than dmg_critical, still under 800ms."),
]
for name, desc in t3_items:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(name)
    r1.bold = True
    r1.font.name = "Consolas"
    p.add_run(f" — {desc}")

# ─── FORMAT SPEC ────────────────────────────────────────────
h2("Format Spec")
bullet("Sample rate: 44.1kHz")
bullet("Bit depth: 16-bit (24-bit OK if you prefer)")
bullet("Format: WAV masters + OGG/Vorbis exports")
bullet("Length: UI cues <300ms, weapon fire/impact <800ms, lifecycle stings 1.5–3s")
bullet("Stems separated for weapon cues (fire layer + impact layer as separate files) so we can mix in-engine")
bullet("Mono unless stereo width matters (Crazy Ivan, Heatseeker can be stereo for spread)")
bullet("Headroom: master peaks at -3dBTP, target loudness ~-14 LUFS for SFX (no aggressive limiter)")
bullet("Naming: snake_case, prefix-grouped — weap_single_fire.wav, weap_single_impact.wav, ui_tap.wav, dmg_critical.wav, match_end_victory.wav")

# ─── CONSTRAINTS ────────────────────────────────────────────
h2("Constraints")
bullet("Mobile speaker compatibility — primary playback surface is a phone speaker through Telegram. Anything below ~100Hz won't reproduce. Mix accordingly; sub-bass is wasted.")
bullet("Polyphony budget — assume up to 8 SFX overlapping (4-player multi-shot scatter weapon during a kill). Cues should mix naturally, not all peak in the same frequency band.")
bullet("No vocal stings.")
bullet("No copyrighted samples — must be original or licensed for game-commercial use, royalty-free in perpetuity, transferable license.")
bullet("Stems delivery for weapons — fire / impact as separate files where applicable.")

# ─── REFERENCES ─────────────────────────────────────────────
h2("Reference Tracks for Tone / Energy")
body("Anchor on the feel, don't have to literally match. Listen to one full clip of each before quoting.")
bullet("iShoot (2008) — weapon variety, retro charm")
bullet("Pocket Tanks (2001) — UI restraint, weapon distinctiveness")
bullet("Worms Armageddon — slightly comedic violence energy")
bullet("Hi-Fi Rush — UI rhythm and tactile clicks")
bullet("Hotline Miami — punchy lo-fi impact philosophy")

# ─── PROCESS ────────────────────────────────────────────────
h2("Process")
bullet("Free sample request before commit — 1 weapon (Single Shot fire+impact) so we can validate tone before ordering the full 27.")
bullet("Round 1 delivery (after sample approved): all 27 cues, 7-day turnaround.")
bullet("Up to 2 rounds of revisions on whatever doesn't land.")

body("")
p = doc.add_paragraph()
r = p.add_run("Total budget target: £300–500 (≈ $400–600). 27 cues, single seller for consistency.")
r.bold = True

# ─── OUTPUT STRUCTURE ───────────────────────────────────────
h2("Delivery Structure")
p = doc.add_paragraph()
r = p.add_run(
    "solshot_audio_custom/\n"
    "├── weapons/\n"
    "│   ├── weap_single_fire.wav\n"
    "│   ├── weap_single_impact.wav\n"
    "│   ├── weap_big_fire.wav\n"
    "│   ├── weap_big_impact.wav\n"
    "│   ├── weap_heatseeker_fire.wav\n"
    "│   ├── weap_heatseeker_inflight.wav\n"
    "│   ├── weap_heatseeker_impact.wav\n"
    "│   └── ... (15 weapon cues total)\n"
    "├── lifecycle/\n"
    "│   ├── match_start_fanfare.wav\n"
    "│   ├── match_end_victory.wav\n"
    "│   ├── match_end_defeat.wav\n"
    "│   ├── you_eliminated.wav\n"
    "│   ├── tank_eliminated.wav\n"
    "│   ├── ui_shot_burn.wav\n"
    "│   └── ui_gold_earned_kill.wav\n"
    "├── ui/\n"
    "│   ├── ui_tap.wav\n"
    "│   ├── ui_modal_open.wav\n"
    "│   ├── ui_modal_close.wav\n"
    "│   ├── ui_toast_error.wav\n"
    "│   ├── dmg_critical.wav\n"
    "│   └── dmg_devastating.wav\n"
    "└── README.md  (mix notes, recommended ducking, polyphony tips)"
)
r.font.name = "Consolas"
r.font.size = Pt(10)

# ─── CLOSING ────────────────────────────────────────────────
h2("Closing Note")
body(
    "This brief is the *identity layer* of SolShot's audio. The remaining ~50-60 cues "
    "(other weapons, generic UI, ambient biome loops, movement) will be sourced from "
    "free libraries (Mixkit, Freesound CC0). Please don't quote for those — focus on "
    "making the 27 cues above feel unmistakably ours."
)
body("Reach out with the free sample first. Looking forward.")

# Save
doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Size:  {OUT.stat().st_size:,} bytes")
